from django.utils.timezone import now
from django.db import models
import PyPDF2
import docx
import win32com.client
import pythoncom 
import os
from django.conf import settings
import hashlib
from django.core.exceptions import ValidationError
class Domaine(models.Model):  # ou Theme
    nom = models.CharField(max_length=100, unique=True,null=True)

    def __str__(self):
        return self.nom
    class Meta:
        ordering = []
# Create your models here.
class Document(models.Model):
    TYPE_CHOICES = [
        ('Constitution', 'Constitution'),
        ('Traités internationaux', 'Traités Internationaux'),
        ('Convention', 'Convention'),
        ('Lois organiques', 'Lois Organiques'),
        ('Lois ordinaires', 'Lois Ordinaires'),
        ('Ordonnances', 'Ordonnances'),
        ('Décrets', 'Décrets'),
        ('Arrêtés interministeriels', 'Arrêtés Interministériels'),
        ('Arrêtés', 'Arrêtés'),
        ('Circilaire', 'Circulaire'),
        ('Notes', 'Notes'),
    ]

    CONSEIL_CHOICES = [
        ('Autre', 'Autre'),
        ('ministre', 'Ministre'),
        ('gouvernement', 'Gouvernement'),
    ]

    STATUS_CHOICES = [
        ('En vigueur', 'En vigueur'),
        ('Abrogé', 'Abrogé'),
    ]

    type = models.CharField(max_length=50, choices=TYPE_CHOICES)
    objet = models.TextField()
    numero = models.TextField()
    date = models.DateField()
    conseil = models.CharField(max_length=50, choices=CONSEIL_CHOICES, default='Autre')
    domaine = models.ForeignKey('Domaine', on_delete=models.CASCADE, null=True,blank=True)
    fichier = models.FileField(upload_to='documents/')
    pdf_file = models.FileField(upload_to='pdf_documents/', null=True, blank=True)
    status = models.CharField(max_length=10, choices=STATUS_CHOICES, default='En vigueur')
    last_modified_by = models.ForeignKey(
        settings.AUTH_USER_MODEL,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='modified_documents'
    )
    last_modified_at = models.DateTimeField(null=True, blank=True)
    modification_details = models.TextField(null=True, blank=True)  # Description des modifications
    inclus_journal = models.BooleanField(default=False)
    date_journal = models.DateField(null=True, blank=True)
    numero_journal = models.CharField(max_length=20, null=True, blank=True)
    page_journal = models.CharField(max_length=20, null=True, blank=True)
    visits = models.PositiveIntegerField(default=0,verbose_name='nombre de visite')
    telechargements = models.PositiveIntegerField(default=0)

    # def increment_visits(self):
    #     self.visits += 1
    #     self.save()
    # def increment_telechargements(self):
    #     self.telechargements += 1
    #     self.save()
    def increment_visits(self):
        Document.objects.filter(id=self.id).update(visits=models.F('visits') + 1)

    def increment_telechargements(self):
        Document.objects.filter(id=self.id).update(telechargements=models.F('telechargements') + 1)
    def update_modification(self, user, details):
        """Mise à jour des informations de modification."""
        self.last_modified_by = user
        self.last_modified_at = now()
        self.modification_details = details
        self.save()

    def save(self, *args, **kwargs):
        # Vérifier si l'objet est en cours de création (pas de modification)
        if not self.pk:  # self.pk est None lors de la création
            if self.fichier:
                try:
                    print(f"Nom du fichier : {self.fichier.name}")  # Log pour vérifier le nom du fichier
                    print(f"Taille du fichier : {self.fichier.size}")  # Log pour vérifier la taille du fichier

                    # Calculer le hash du fichier
                    file_content = self.fichier.read()
                    file_hash = hashlib.sha256(file_content).hexdigest()
                    print(f"Hash calculé : {file_hash}")  # Log pour vérifier le hash

                    # Vérifier si un fichier avec le même hash existe déjà
                    for existing_doc in Document.objects.all():
                        try:
                            existing_file_content = existing_doc.fichier.read()
                            existing_file_hash = hashlib.sha256(existing_file_content).hexdigest()
                            print(f"Hash du fichier existant ({existing_doc.fichier.name}) : {existing_file_hash}")  # Log pour vérifier les hashs existants
                            if file_hash == existing_file_hash:
                                raise ValidationError("Un fichier identique existe déjà.")
                            existing_doc.fichier.seek(0)  # Réinitialiser le fichier après la lecture
                        except FileNotFoundError:
                            print(f"Fichier manquant : {existing_doc.fichier.name}")  # Log pour les fichiers manquants
                            continue  # Ignorer les fichiers manquants

                    # Réinitialiser le fichier après la lecture
                    self.fichier.seek(0)

                except Exception as e:
                    print(f"Erreur lors de la vérification du fichier : {str(e)}")  # Log pour les erreurs
                    raise ValidationError(f"Erreur lors de la vérification du fichier : {str(e)}")

        # Appeler la méthode save() de la classe parente
        super().save(*args, **kwargs)
    def __str__(self):
        return f"{self.type} - {self.numero}"
    
    def convert_to_pdf(self):
        if not self.fichier.name.endswith('.docx'):
            raise ValueError("La conversion en PDF est uniquement prise en charge pour les fichiers .docx.")

        input_path = self.fichier.path
        media_root = os.path.dirname(os.path.dirname(input_path))
        output_dir = os.path.join(media_root, 'pdf_documents')
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, os.path.basename(input_path).replace('.docx', '.pdf'))

        try:
            pythoncom.CoInitialize()
            word = win32com.client.Dispatch("Word.Application")
            doc = word.Documents.Open(input_path)
            doc.SaveAs(output_path, FileFormat=17)
            doc.Close()
            word.Quit()

            self.pdf_file.name = os.path.relpath(output_path, media_root)
            self.save()

        except Exception as e:
            raise ValueError(f"Erreur lors de la conversion : {e}")


        return output_path  # Retourne le chemin absolu du fichier PDF


        
    def extract_content(self):
        if self.fichier.name.endswith('.pdf'):
            try:
                with self.fichier.open('rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    return ''.join(page.extract_text() for page in reader.pages)
            except Exception as e:
                return f"Error reading PDF: {e}"
        elif self.fichier.name.endswith(('.doc', '.docx')):
            try:
                with self.fichier.open('rb') as f:
                    doc = docx.Document(f)
                    return '\n'.join(p.text for p in doc.paragraphs)
            except Exception as e:
                return f"Error reading DOC/DOCX: {e}"
        return "Unsupported file type."
    class Meta:
        indexes = [
            models.Index(fields=['objet'],name='objet_idx'),
            models.Index(fields=['numero'],name='numero_idx'),
        ]

class DocumentStats(models.Model):
    date = models.DateField(default=now)  # Date de l'ajout
    daily_count = models.IntegerField(default=0)  # Nombre d'ajouts par jour
    monthly_count = models.IntegerField(default=0)  # Nombre d'ajouts par mois
    yearly_count = models.IntegerField(default=0)  # Nombre d'ajouts par année

    class Meta:
        verbose_name_plural = "Document Statistics"
class Actualite(models.Model):
    CONSEIL_CHOICES = [
        ('CONSEIL DES MINISTRES', 'Ministre'),
        ('CONSEIL DU GOUVERNEMENT', 'Gouvernement'),
    ]
    conseil = models.CharField(max_length=50,choices=CONSEIL_CHOICES,default='CONSEIL DES MINISTRES')
    titre = models.CharField(max_length=200)
    date = models.DateField(default='2024-01-01')
    lieu = models.CharField(max_length=100)
    texte = models.TextField()

    def __str__(self):
        return self.titre

class Remark(models.Model):
    email = models.EmailField()
    message = models.TextField()
    created_at = models.DateTimeField(auto_now_add=True)

    def __str__(self):
        return self.email

